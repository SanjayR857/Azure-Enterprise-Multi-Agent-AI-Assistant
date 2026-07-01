import pandas as pd
# pyrefly: ignore [missing-import]
from pypdf import PdfReader

# pyrefly: ignore [missing-import]
from docx import Document

from langchain_core.documents import Document as LangchainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
# pyrefly: ignore [missing-import]
from langchain_chroma import Chroma
from langchain_ollama import OllamaEmbeddings
import logging

logger = logging.getLogger(__name__)

# Initialize the Embeddings Model
# NOTE: Make sure you have pulled an embeddings model in Ollama (e.g., run `ollama pull nomic-embed-text` in your terminal)
embeddings = OllamaEmbeddings(model="nomic-embed-text") 

# Initialize ChromaDB Vector Store
vector_store = Chroma(
    collection_name="chat_documents",
    embedding_function=embeddings,
    persist_directory="./chroma_db"
)

class DocumentParserService:
    def extract_text(self, file_path: str, ext: str) -> str:
        logger.info("Starting text extraction", extra={"file_path": file_path, "extension": ext})
        text = ""
        try:
            if ext == ".pdf":
                reader = PdfReader(file_path)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            elif ext == ".docx":
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif ext in [".xlsx", ".xls"]:
                df = pd.read_excel(file_path)
                text = df.to_string()
            elif ext == ".csv":
                df = pd.read_csv(file_path)
                text = df.to_string()
            elif ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            else:
                logger.warning("Unsupported extraction", extra={"extension": ext, "file_path": file_path})
        except Exception as e:
            logger.error("Error parsing document", extra={"file_path": file_path, "error": str(e)})
            
        logger.info("Finished extracting text", extra={"file_path": file_path, "text_length": len(text)})
        return text

    async def process_and_index(self, db_document: dict):
        """
        Extracts text from the physical file, chunks it, and saves to Vector DB.
        """
        original_filename = db_document["original_filename"]
        logger.info("Starting process_and_index for document", extra={"original_filename": original_filename})
        # Get the extension
        ext = original_filename[original_filename.rfind("."):].lower()
        
        # 1. Extract raw text from the file
        raw_text = self.extract_text(db_document["storage_path"], ext)
        if not raw_text.strip():
            logger.warning("No text extracted", extra={"original_filename": original_filename})
            return

        # 2. Chop the text into ~1000 character chunks with a 200 character overlap
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            add_start_index=True
        )
        chunks = text_splitter.split_text(raw_text)

        # 3. Create Langchain Documents and attach our Database Metadata to them
        docs = []
        for chunk in chunks:
            docs.append(LangchainDocument(
                page_content=chunk,
                metadata={
                    "session_id": db_document["session_id"],
                    "document_id": db_document["id"],
                    "filename": original_filename
                }
            ))

        # 4. Embed and save to Chroma DB
        if docs:
            vector_store.add_documents(docs)
            logger.info("Indexed chunks for document", extra={"original_filename": db_document['original_filename'], "chunk_count": len(docs)})

document_parser = DocumentParserService()
